from langchain_core.tools import tool
import os
from dotenv import load_dotenv
from langchain_google_genai import ChatGoogleGenerativeAI
from google.oauth2.service_account import Credentials
from google.oauth2.credentials import Credentials as OAuth2Credentials
from google_auth_oauthlib.flow import InstalledAppFlow
from google.auth.transport.requests import Request
from googleapiclient.discovery import build
from googleapiclient.http import MediaIoBaseDownload
import io
import PyPDF2
import docx
import json
import base64
from email import message_from_bytes
from datetime import datetime
import re
import pickle


def get_gmail_credentials_oauth():
    """Get Gmail credentials using OAuth2 for personal Gmail accounts."""
    SCOPES = [
        'https://www.googleapis.com/auth/gmail.readonly',
        'https://www.googleapis.com/auth/gmail.send'
    ]
    creds = None

    # Token file stores the user's access and refresh tokens
    token_file = 'token.pickle'

    # Check if token file exists
    if os.path.exists(token_file):
        with open(token_file, 'rb') as token:
            creds = pickle.load(token)

    # If there are no valid credentials, let the user log in
    if not creds or not creds.valid:
        if creds and creds.expired and creds.refresh_token:
            creds.refresh(Request())
        else:
            flow = InstalledAppFlow.from_client_secrets_file(
                'credentials.json', SCOPES)
            creds = flow.run_local_server(port=0)

        # Save the credentials for the next run
        with open(token_file, 'wb') as token:
            pickle.dump(creds, token)

    return creds


def create_cv_folder():
    """Create CV collection folder if it doesn't exist."""
    folder_path = os.path.join(os.getcwd(), "cv_collection")
    if not os.path.exists(folder_path):
        os.makedirs(folder_path)
    return folder_path


def sanitize_filename(filename):
    """Sanitize filename to remove invalid characters."""
    # Remove invalid characters for filenames
    filename = re.sub(r'[<>:"/\\|?*]', '_', filename)
    return filename


def download_cv_from_gmail(gmail_service, cv_folder):
    """Download CV attachments from Gmail and save to local folder."""
    downloaded_cvs = []

    # Load processed message IDs to prevent re-downloading from same email
    processed_ids_file = os.path.join(cv_folder, '.processed_message_ids.txt')
    processed_message_ids = set()
    if os.path.exists(processed_ids_file):
        with open(processed_ids_file, 'r') as f:
            processed_message_ids = set(line.strip() for line in f)

    try:
        # Search for emails with CV attachments (last 30 days)
        # Simplified query: just look for PDF/DOC/DOCX attachments without subject filtering
        query = 'has:attachment (filename:pdf OR filename:doc OR filename:docx)'
        results = gmail_service.users().messages().list(
            userId='me',
            q=query,
            maxResults=50
        ).execute()

        messages = results.get('messages', [])

        if not messages:
            return downloaded_cvs

        for message in messages:
            msg_id = message['id']

            # Skip if we've already processed this email message
            if msg_id in processed_message_ids:
                continue

            msg = gmail_service.users().messages().get(
                userId='me',
                id=msg_id
            ).execute()

            # Get email metadata
            headers = msg['payload']['headers']
            subject = next((h['value'] for h in headers if h['name'] == 'Subject'), 'No Subject')
            sender = next((h['value'] for h in headers if h['name'] == 'From'), 'Unknown')
            date = next((h['value'] for h in headers if h['name'] == 'Date'), '')

            # Check for attachments
            if 'parts' in msg['payload']:
                for part in msg['payload']['parts']:
                    if part.get('filename'):
                        filename = part['filename']

                        # Only process CV-related files
                        if filename.lower().endswith(('.pdf', '.doc', '.docx')):
                            if 'body' in part and 'attachmentId' in part['body']:
                                attachment_id = part['body']['attachmentId']

                                # Download attachment
                                attachment = gmail_service.users().messages().attachments().get(
                                    userId='me',
                                    messageId=msg_id,
                                    id=attachment_id
                                ).execute()

                                file_data = base64.urlsafe_b64decode(attachment['data'])

                                # Generate unique filename with timestamp
                                safe_filename = sanitize_filename(filename)
                                timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                                unique_filename = f"{timestamp}_{safe_filename}"
                                filepath = os.path.join(cv_folder, unique_filename)

                                # Save file
                                with open(filepath, 'wb') as f:
                                    f.write(file_data)

                                downloaded_cvs.append({
                                    'filename': unique_filename,
                                    'original_filename': filename,
                                    'filepath': filepath,
                                    'sender': sender,
                                    'subject': subject,
                                    'date': date,
                                    'message_id': msg_id  # Add message ID for tracking
                                })

            # Mark this message as processed
            processed_message_ids.add(msg_id)
            with open(processed_ids_file, 'a') as f:
                f.write(f"{msg_id}\n")

        return downloaded_cvs

    except Exception as e:
        print(f"Error downloading CVs from Gmail: {str(e)}")
        return downloaded_cvs


def extract_text_from_pdf(file_content):
    """Extract text from PDF file content."""
    try:
        pdf_file = io.BytesIO(file_content)
        pdf_reader = PyPDF2.PdfReader(pdf_file)
        text = ""
        for page in pdf_reader.pages:
            text += page.extract_text()
        return text
    except Exception as e:
        return f"Error extracting PDF: {str(e)}"


def extract_text_from_pdf_file(filepath):
    """Extract text from PDF file path."""
    try:
        with open(filepath, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            text = ""
            for page in pdf_reader.pages:
                text += page.extract_text()
        return text
    except Exception as e:
        return f"Error extracting PDF: {str(e)}"


def extract_text_from_docx(file_content):
    """Extract text from DOCX file content."""
    try:
        docx_file = io.BytesIO(file_content)
        doc = docx.Document(docx_file)
        text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
        return text
    except Exception as e:
        return f"Error extracting DOCX: {str(e)}"


def extract_text_from_docx_file(filepath):
    """Extract text from DOCX file path."""
    try:
        doc = docx.Document(filepath)
        text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
        return text
    except Exception as e:
        return f"Error extracting DOCX: {str(e)}"


def save_individual_cv_analysis(cv_info, analysis, cv_folder):
    """Save individual ATS analysis result for each CV as a JSON file."""
    try:
        # Create the analysis filename based on the CV filename
        cv_filename = cv_info.get('filename', 'unknown')
        base_name = os.path.splitext(cv_filename)[0]
        analysis_filename = f"{base_name}_ATS_Analysis.json"
        analysis_filepath = os.path.join(cv_folder, analysis_filename)

        # Prepare the detailed analysis report
        detailed_report = {
            "cv_filename": cv_info.get('filename', 'Unknown'),
            "original_filename": cv_info.get('original_filename', 'Unknown'),
            "sender": cv_info.get('sender', 'Unknown'),
            "email_subject": cv_info.get('subject', 'N/A'),
            "email_date": cv_info.get('date', 'N/A'),
            "analysis_date": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
            "ats_score": analysis.get("ats_score", 0),
            "rank": cv_info.get('rank', 0),
            "strengths": analysis.get("strengths", []),
            "improvements": analysis.get("improvements", []),
            "keywords": analysis.get("keywords", []),
            "assessment": analysis.get("assessment", "")
        }

        # Save to JSON file
        with open(analysis_filepath, 'w', encoding='utf-8') as f:
            json.dump(detailed_report, f, indent=2, ensure_ascii=False)

        return analysis_filepath
    except Exception as e:
        print(f"Error saving analysis for {cv_info.get('filename', 'unknown')}: {str(e)}")
        return None


def analyze_resume_ats(resume_text, llm, job_description=""):
    """Analyze resume for ATS compatibility and score it."""

    job_context = f"\n\nJob Description Context:\n{job_description}\n" if job_description else ""

    prompt = f"""
    You are an expert ATS (Applicant Tracking System) analyzer and HR professional.
    Analyze the following resume and provide:

    1. ATS Score (0-100): How well this resume would perform in an ATS system
    2. Key Strengths (3-5 points)
    3. Areas for Improvement (3-5 points)
    4. Keywords Found (list relevant industry keywords)
    5. Overall Assessment (2-3 sentences)
    {job_context}
    Consider these ATS factors:
    - Formatting and structure
    - Use of relevant keywords
    - Clear section headings
    - Quantifiable achievements
    - Skills alignment
    - Professional experience clarity
    - Match with job requirements (if job description provided)

    Resume Text:
    {resume_text[:4000]}

    Provide your response in JSON format with keys: ats_score, strengths, improvements, keywords, assessment
    """

    try:
        response = llm.invoke(prompt)
        # Extract JSON from response
        content = response.content if hasattr(response, 'content') else str(response)

        # Try to parse JSON from the response
        if "```json" in content:
            json_str = content.split("```json")[1].split("```")[0].strip()
        elif "```" in content:
            json_str = content.split("```")[1].split("```")[0].strip()
        else:
            json_str = content

        analysis = json.loads(json_str)
        return analysis
    except Exception as e:
        # Fallback: create a basic score
        return {
            "ats_score": 50,
            "strengths": ["Resume received"],
            "improvements": ["Analysis error occurred"],
            "keywords": [],
            "assessment": f"Error during analysis: {str(e)}"
        }


@tool
def collect_and_review_cvs_from_gmail(job_description: str = "", search_query: str = ""):
    """
    Automatically downloads CV attachments from Gmail to a local 'cv_collection' folder,
    analyzes ATS rates, scores each resume, and returns the top 10 ranked candidates.

    Uses OAuth2 authentication for personal Gmail accounts.

    Args:
        job_description (str): Optional job description to match against (for better analysis)
        search_query (str): Optional custom Gmail search query (default searches for CVs/resumes)

    Returns:
        str: JSON string containing download summary, top 10 ranked resumes with their scores and analysis
    """
    try:
        load_dotenv()

        # Create CV collection folder
        cv_folder = create_cv_folder()

        # Get OAuth2 credentials
        creds = get_gmail_credentials_oauth()

        gmail_service = build('gmail', 'v1', credentials=creds)

        # Download CVs from Gmail
        print(f"Downloading CVs from Gmail to {cv_folder}...")
        downloaded_cvs = download_cv_from_gmail(gmail_service, cv_folder)

        if not downloaded_cvs:
            return json.dumps({
                "error": "No CV attachments found in Gmail",
                "cv_folder": cv_folder,
                "downloaded_count": 0,
                "top_resumes": []
            })

        # AUTOMATICALLY send acknowledgment emails to each CV sender
        print(f"Sending acknowledgment emails to {len(downloaded_cvs)} CV senders...")
        from .email_tools import log_cv_sender_to_csv, get_gmail_service, create_email_message

        email_service = get_gmail_service()
        sender_email = "hr.agent.automation@gmail.com"

        for cv_info in downloaded_cvs:
            try:
                # Extract sender email and name
                sender_full = cv_info['sender']
                recipient_email = sender_full.split('<')[-1].replace('>', '').strip() if '<' in sender_full else sender_full
                recipient_name = sender_full.split('<')[0].strip() if '<' in sender_full else "Applicant"

                # Send acknowledgment email
                subject = "Thank You for Your Application - CV Received"
                body = f"""Dear {recipient_name},

Thank you for submitting your CV/Resume to our team.

We have successfully received your application and it is currently being reviewed by our HR department. Our team will carefully assess your qualifications and experience.

We will contact you as soon as possible regarding the next steps in the recruitment process. This typically takes 3-5 business days.

If you have any questions in the meantime, please feel free to reach out to us.

Thank you for your interest in joining our organization.

Best regards,
HR Department
{sender_email}

---
This is an automated acknowledgment email. Please do not reply to this message."""

                email_message = create_email_message(sender_email, recipient_email, subject, body)
                email_service.users().messages().send(userId='me', body=email_message).execute()

                # Log to CSV
                log_cv_sender_to_csv({
                    'name': recipient_name,
                    'email': recipient_email,
                    'cv_filename': cv_info['filename'],
                    'date_sent': cv_info['date'],
                    'reply_sent': 'Yes'
                }, cv_folder)

                print(f"  ✓ Sent acknowledgment to {recipient_email}")

            except Exception as email_error:
                print(f"  ✗ Failed to send acknowledgment to {cv_info['sender']}: {str(email_error)}")
                # Log failed attempt
                log_cv_sender_to_csv({
                    'name': recipient_name if 'recipient_name' in locals() else 'Unknown',
                    'email': recipient_email if 'recipient_email' in locals() else cv_info['sender'],
                    'cv_filename': cv_info['filename'],
                    'date_sent': cv_info['date'],
                    'reply_sent': 'No'
                }, cv_folder)

        print(f"✓ Acknowledgment emails sent complete")

        # Check for Google API key
        google_api_key = os.getenv("GOOGLE_API_KEY")
        if not google_api_key or google_api_key == "YOUR_GOOGLE_GEMINI_API_KEY_HERE":
            return json.dumps({
                "error": "GOOGLE_API_KEY is not configured. Please add your Google Gemini API key to the .env file.",
                "instructions": "Get your API key from https://makersuite.google.com/app/apikey and add it to backend/.env as GOOGLE_API_KEY='your-key-here'",
                "downloaded_count": len(downloaded_cvs),
                "top_resumes": []
            })

        # Initialize LLM for analysis
        llm = ChatGoogleGenerativeAI(model="gemini-1.5-flash", temperature=0.3)

        resume_analyses = []

        # Process each downloaded CV
        for cv_info in downloaded_cvs:
            try:
                filepath = cv_info['filepath']
                filename = cv_info['filename']

                # Extract text based on file type
                if filepath.lower().endswith('.pdf'):
                    resume_text = extract_text_from_pdf_file(filepath)
                elif filepath.lower().endswith(('.doc', '.docx')):
                    resume_text = extract_text_from_docx_file(filepath)
                else:
                    continue

                # Analyze resume
                analysis = analyze_resume_ats(resume_text, llm, job_description)

                cv_analysis_data = {
                    "rank": 0,  # Will be assigned after sorting
                    "file_name": filename,
                    "original_filename": cv_info['original_filename'],
                    "filepath": filepath,
                    "sender": cv_info['sender'],
                    "email_subject": cv_info['subject'],
                    "email_date": cv_info['date'],
                    "ats_score": analysis.get("ats_score", 0),
                    "strengths": analysis.get("strengths", []),
                    "improvements": analysis.get("improvements", []),
                    "keywords": analysis.get("keywords", []),
                    "assessment": analysis.get("assessment", "")
                }

                resume_analyses.append(cv_analysis_data)

            except Exception as file_error:
                # Log error but continue with other files
                resume_analyses.append({
                    "rank": 0,
                    "file_name": cv_info.get('filename', 'Unknown'),
                    "filepath": cv_info.get('filepath', ''),
                    "sender": cv_info.get('sender', 'Unknown'),
                    "ats_score": 0,
                    "strengths": [],
                    "improvements": [],
                    "keywords": [],
                    "assessment": f"Error processing file: {str(file_error)}"
                })

        # Sort by ATS score (highest first)
        resume_analyses.sort(key=lambda x: x['ats_score'], reverse=True)

        # Assign ranks to all resumes and save individual analysis files
        saved_analysis_files = []
        for rank, resume in enumerate(resume_analyses, 1):
            resume['rank'] = rank

            # Save individual ATS analysis file for each CV
            cv_info_for_save = {
                'filename': resume.get('file_name'),
                'original_filename': resume.get('original_filename'),
                'sender': resume.get('sender'),
                'subject': resume.get('email_subject'),
                'date': resume.get('email_date'),
                'rank': rank
            }
            analysis_for_save = {
                'ats_score': resume.get('ats_score'),
                'strengths': resume.get('strengths'),
                'improvements': resume.get('improvements'),
                'keywords': resume.get('keywords'),
                'assessment': resume.get('assessment')
            }

            analysis_file = save_individual_cv_analysis(cv_info_for_save, analysis_for_save, cv_folder)
            if analysis_file:
                saved_analysis_files.append(analysis_file)

        result = {
            "cv_folder": cv_folder,
            "downloaded_count": len(downloaded_cvs),
            "total_resumes_analyzed": len(resume_analyses),
            "analysis_files_created": len(saved_analysis_files),
            "all_resumes": resume_analyses,  # Return all individual results
            "top_10_resumes": resume_analyses[:10],  # Also provide top 10 for convenience
            "summary": f"Downloaded {len(downloaded_cvs)} CVs from Gmail. Analyzed {len(resume_analyses)} resumes. Created {len(saved_analysis_files)} individual ATS analysis files. Top candidate: {resume_analyses[0]['original_filename'] if resume_analyses else 'N/A'} with ATS score: {resume_analyses[0]['ats_score'] if resume_analyses else 0}"
        }

        return json.dumps(result, indent=2)

    except Exception as e:
        return json.dumps({
            "error": f"Error collecting and reviewing CVs: {str(e)}",
            "top_resumes": []
        })


@tool
def review_existing_cvs_in_folder(job_description: str = ""):
    """
    Reviews CVs that are already downloaded in the local 'cv_collection' folder,
    analyzes ATS rates, scores each resume, creates individual analysis files,
    and returns all ranked candidates.

    Args:
        job_description (str): Optional job description to match against (for better analysis)

    Returns:
        str: JSON string containing all ranked resumes with their scores and analysis
    """
    try:
        load_dotenv()

        # Get CV collection folder
        cv_folder = create_cv_folder()

        # Check if folder exists and has files
        if not os.path.exists(cv_folder):
            return json.dumps({
                "error": "cv_collection folder does not exist",
                "cv_folder": cv_folder,
                "total_resumes_analyzed": 0,
                "all_resumes": []
            })

        # Get all CV files in the folder (skip JSON analysis files)
        all_files = os.listdir(cv_folder)
        cv_files = [f for f in all_files if f.lower().endswith(('.pdf', '.doc', '.docx'))]

        if not cv_files:
            return json.dumps({
                "error": "No CV files found in cv_collection folder",
                "cv_folder": cv_folder,
                "total_resumes_analyzed": 0,
                "all_resumes": []
            })

        # Check for Google API key
        google_api_key = os.getenv("GOOGLE_API_KEY")
        if not google_api_key or google_api_key == "YOUR_GOOGLE_GEMINI_API_KEY_HERE":
            return json.dumps({
                "error": "GOOGLE_API_KEY is not configured. Please add your Google Gemini API key to the .env file.",
                "instructions": "Get your API key from https://makersuite.google.com/app/apikey and add it to backend/.env as GOOGLE_API_KEY='your-key-here'",
                "cv_folder": cv_folder,
                "total_resumes_analyzed": 0,
                "all_resumes": []
            })

        # Initialize LLM for analysis
        llm = ChatGoogleGenerativeAI(model="gemini-1.5-flash", temperature=0.3)

        resume_analyses = []

        # Process each CV file
        for cv_file in cv_files:
            try:
                filepath = os.path.join(cv_folder, cv_file)

                # Extract text based on file type
                if filepath.lower().endswith('.pdf'):
                    resume_text = extract_text_from_pdf_file(filepath)
                elif filepath.lower().endswith(('.doc', '.docx')):
                    resume_text = extract_text_from_docx_file(filepath)
                else:
                    continue

                # Analyze resume
                analysis = analyze_resume_ats(resume_text, llm, job_description)

                cv_analysis_data = {
                    "rank": 0,  # Will be assigned after sorting
                    "file_name": cv_file,
                    "filepath": filepath,
                    "ats_score": analysis.get("ats_score", 0),
                    "strengths": analysis.get("strengths", []),
                    "improvements": analysis.get("improvements", []),
                    "keywords": analysis.get("keywords", []),
                    "assessment": analysis.get("assessment", "")
                }

                resume_analyses.append(cv_analysis_data)

            except Exception as file_error:
                # Log error but continue with other files
                resume_analyses.append({
                    "rank": 0,
                    "file_name": cv_file,
                    "filepath": os.path.join(cv_folder, cv_file),
                    "ats_score": 0,
                    "strengths": [],
                    "improvements": [],
                    "keywords": [],
                    "assessment": f"Error processing file: {str(file_error)}"
                })

        # Sort by ATS score (highest first)
        resume_analyses.sort(key=lambda x: x['ats_score'], reverse=True)

        # Assign ranks to all resumes and save individual analysis files
        saved_analysis_files = []
        for rank, resume in enumerate(resume_analyses, 1):
            resume['rank'] = rank

            # Save individual ATS analysis file for each CV
            cv_info_for_save = {
                'filename': resume.get('file_name'),
                'original_filename': resume.get('file_name'),
                'sender': 'Local Folder',
                'subject': 'Previously Downloaded CV',
                'date': 'N/A',
                'rank': rank
            }
            analysis_for_save = {
                'ats_score': resume.get('ats_score'),
                'strengths': resume.get('strengths'),
                'improvements': resume.get('improvements'),
                'keywords': resume.get('keywords'),
                'assessment': resume.get('assessment')
            }

            analysis_file = save_individual_cv_analysis(cv_info_for_save, analysis_for_save, cv_folder)
            if analysis_file:
                saved_analysis_files.append(analysis_file)

        result = {
            "cv_folder": cv_folder,
            "total_resumes_analyzed": len(resume_analyses),
            "analysis_files_created": len(saved_analysis_files),
            "all_resumes": resume_analyses,
            "top_10_resumes": resume_analyses[:10],
            "summary": f"Analyzed {len(resume_analyses)} CVs from local folder. Created {len(saved_analysis_files)} individual ATS analysis files. Top candidate: {resume_analyses[0]['file_name'] if resume_analyses else 'N/A'} with ATS score: {resume_analyses[0]['ats_score'] if resume_analyses else 0}"
        }

        return json.dumps(result, indent=2)

    except Exception as e:
        return json.dumps({
            "error": f"Error reviewing existing CVs: {str(e)}",
            "all_resumes": []
        })


@tool
def review_resumes_from_drive(folder_id: str, job_description: str = ""):
    """
    Reviews resumes from a Google Drive folder, analyzes ATS rates, scores each resume,
    and returns the top 10 ranked candidates.

    Args:
        folder_id (str): The Google Drive folder ID containing the resumes
        job_description (str): Optional job description to match against (for better analysis)

    Returns:
        str: JSON string containing top 10 ranked resumes with their scores and analysis
    """
    try:
        load_dotenv()

        # Initialize Google Drive API
        scope = [
            "https://www.googleapis.com/auth/drive.readonly",
            "https://www.googleapis.com/auth/drive.metadata.readonly"
        ]

        creds = Credentials.from_service_account_file(
            "neon-poetry-467907-k3-c6c2f44fb294.json",
            scopes=scope
        )

        drive_service = build('drive', 'v3', credentials=creds)

        # Check for Google API key
        google_api_key = os.getenv("GOOGLE_API_KEY")
        if not google_api_key or google_api_key == "YOUR_GOOGLE_GEMINI_API_KEY_HERE":
            return json.dumps({
                "error": "GOOGLE_API_KEY is not configured. Please add your Google Gemini API key to the .env file.",
                "instructions": "Get your API key from https://makersuite.google.com/app/apikey and add it to backend/.env as GOOGLE_API_KEY='your-key-here'",
                "top_resumes": []
            })

        # Initialize LLM for analysis
        llm = ChatGoogleGenerativeAI(model="gemini-1.5-flash", temperature=0.3)

        # Query files in the folder
        query = f"'{folder_id}' in parents and (mimeType='application/pdf' or mimeType='application/vnd.openxmlformats-officedocument.wordprocessingml.document')"
        results = drive_service.files().list(
            q=query,
            fields="files(id, name, mimeType)",
            pageSize=100
        ).execute()

        files = results.get('files', [])

        if not files:
            return json.dumps({"error": "No resume files found in the specified folder", "top_resumes": []})

        resume_analyses = []

        # Process each resume
        for idx, file in enumerate(files, 1):
            try:
                file_id = file['id']
                file_name = file['name']
                mime_type = file['mimeType']

                # Download file content
                request = drive_service.files().get_media(fileId=file_id)
                file_content = io.BytesIO()
                downloader = MediaIoBaseDownload(file_content, request)

                done = False
                while not done:
                    status, done = downloader.next_chunk()

                file_content.seek(0)
                content_bytes = file_content.read()

                # Extract text based on file type
                if mime_type == 'application/pdf':
                    resume_text = extract_text_from_pdf(content_bytes)
                elif mime_type == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document':
                    resume_text = extract_text_from_docx(content_bytes)
                else:
                    continue

                # Analyze resume
                analysis = analyze_resume_ats(resume_text, llm, job_description)

                resume_analyses.append({
                    "rank": 0,  # Will be assigned after sorting
                    "file_name": file_name,
                    "file_id": file_id,
                    "ats_score": analysis.get("ats_score", 0),
                    "strengths": analysis.get("strengths", []),
                    "improvements": analysis.get("improvements", []),
                    "keywords": analysis.get("keywords", []),
                    "assessment": analysis.get("assessment", "")
                })

            except Exception as file_error:
                # Log error but continue with other files
                resume_analyses.append({
                    "rank": 0,
                    "file_name": file.get('name', 'Unknown'),
                    "file_id": file.get('id', ''),
                    "ats_score": 0,
                    "strengths": [],
                    "improvements": [],
                    "keywords": [],
                    "assessment": f"Error processing file: {str(file_error)}"
                })

        # Sort by ATS score (highest first)
        resume_analyses.sort(key=lambda x: x['ats_score'], reverse=True)

        # Create a local folder for Drive CV analyses (optional)
        cv_folder = create_cv_folder()

        # Assign ranks to all resumes and save individual analysis files
        saved_analysis_files = []
        for rank, resume in enumerate(resume_analyses, 1):
            resume['rank'] = rank

            # Save individual ATS analysis file for each CV from Drive
            cv_info_for_save = {
                'filename': resume.get('file_name'),
                'original_filename': resume.get('file_name'),
                'sender': 'Google Drive',
                'subject': f"Drive File ID: {resume.get('file_id')}",
                'date': 'N/A',
                'rank': rank
            }
            analysis_for_save = {
                'ats_score': resume.get('ats_score'),
                'strengths': resume.get('strengths'),
                'improvements': resume.get('improvements'),
                'keywords': resume.get('keywords'),
                'assessment': resume.get('assessment')
            }

            analysis_file = save_individual_cv_analysis(cv_info_for_save, analysis_for_save, cv_folder)
            if analysis_file:
                saved_analysis_files.append(analysis_file)

        result = {
            "cv_folder": cv_folder,
            "total_resumes_analyzed": len(resume_analyses),
            "analysis_files_created": len(saved_analysis_files),
            "all_resumes": resume_analyses,  # Return all individual results
            "top_10_resumes": resume_analyses[:10],  # Also provide top 10 for convenience
            "summary": f"Analyzed {len(resume_analyses)} resumes from Drive. Created {len(saved_analysis_files)} individual ATS analysis files in {cv_folder}. Top candidate: {resume_analyses[0]['file_name'] if resume_analyses else 'N/A'} with ATS score: {resume_analyses[0]['ats_score'] if resume_analyses else 0}"
        }

        return json.dumps(result, indent=2)

    except Exception as e:
        return json.dumps({
            "error": f"Error reviewing resumes: {str(e)}",
            "top_resumes": []
        })
